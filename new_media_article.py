#從selenium引入webdriver
from selenium import webdriver
from selenium.webdriver.chrome.options import Options

#引入beautifulsoup & 時間套件
from bs4 import BeautifulSoup
import time

import os
import pymysql.cursors
import urllib.request #開啟iframe
import ssl
ssl._create_default_https_context = ssl._create_unverified_context #ssl不報錯

connection = pymysql.connect(host='localhost',
                             user='root',
                             password='admin1234',
                             db='new_media',
                             charset='utf8mb4',
                             cursorclass=pymysql.cursors.DictCursor)

from datetime import date
today = str(date.today())
from docx import Document #引入生成word檔套件

#初始化選項
options = Options()
#關掉圖片，可加快爬網頁速度
options.experimental_options['prefs'] = {'profile.default_content_settings':{'images':2},
                                        'profile.default_managed_default_content_settings':{'images':2}}
#產生word檔
document = Document()
document.add_heading(today + '新媒體文章抓取',0)
table = document.add_table(rows=1,cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '標題'
hdr_cells[1].text = '分享數'
hdr_cells[2].text = 'Hashtag'
hdr_cells[3].text = '網站'



#背景執行chrome
driver = webdriver.Chrome(os.getcwd()+"/chromedriver", chrome_options=options)

try:
    with connection.cursor() as cursor:
        
        for i in range(1,2):
            driver.get('https://www.inside.com.tw/?page='+ str(i))
            #轉譯成python看得懂的
            sourceCode = BeautifulSoup(driver.page_source)
            article_box = sourceCode.select('div.post_list-list_style')[0]
            articles = article_box.select('div.post_list_item ')
            for article in articles:
                title = article.select('h3.post_title')[0].text
                date = article.select('li.post_date')[0].text.strip().replace('/','-')
                tags = article.select('a.hero_slide_tag')
                tag_string = ''
                for tag in tags:
                    tag_string += tag.text + ', '
                if date == today:
                    print(title)
                    row_cells = table.add_row().cells
                    row_cells[0].text = title
                    row_cells[1].text = ''
                    row_cells[2].text = tag_string
                    row_cells[3].text = 'inside'
                    
#                     sql = '''
#                         INSERT INTO `new_media`.`articles`(`brand`,`title`,`date`,`tags`) 
#                         VALUES ('{}','{}','{}','{}')
#                         '''.format('inside',title,date,tag_string)
#                     #print(sql)
#                     cursor.execute(sql)
#                     connection.commit()
            
            driver.get('https://technews.tw/page/'+ str(i)+'/')
            sourceCode = BeautifulSoup(driver.page_source)
            article_box = sourceCode.select('div#content')[0]
            articles = article_box.select('header.entry-header')
            for article in articles:
                title = article.select('h1.entry-title')[0].text
                date = article.select('span.body')[1].text.strip().replace(' 年 ','-').replace(' 月 ','-')
                #tags = article.select('a[href*="category"]')
                tags = article.select('span.body')[2].select('a')
                fb_iframe = article.select('iframe')[1]
                response = urllib.request.urlopen(fb_iframe.attrs['src'])
                iframe_soup = BeautifulSoup(response)
                share = iframe_soup.select('span#u_0_2')[0].text
                tag_string = ''
                for tag in tags:
                    tag_string += tag.text + ','
                if date[0:10] == today:
                    #print(date)
                    #print(tag_string)
                    #print(share)
                    row_cells = table.add_row().cells
                    row_cells[0].text = title
                    row_cells[1].text = share
                    row_cells[2].text = tag_string
                    row_cells[3].text = 'TechNews'
                    
#                     sql = '''
#                         INSERT INTO `new_media`.`articles`(`brand`,`title`,`date`,`tags`,`share`) 
#                         VALUES ('{}','{}','{}','{}','{}')
#                         '''.format('TechNews',title,date,tag_string,share)
#                     #print(sql)
#                     cursor.execute(sql)
#                     connection.commit()


        driver.get('https://buzzorange.com/techorange/')
        for i in range(1,2):
            driver.execute_script('window.scrollTo(0,document.body.scrollHeight);')
            time.sleep(2)
        sourceCode = BeautifulSoup(driver.page_source)
        article_box = sourceCode.select('main#main')[0]
        #print(article_box)
        articles = article_box.select('article.post')
        #print(articles)
        for article in articles:
            title = article.select('h4.entry-title')[0].text
            date = article.select('time.updated')[0].text.strip().replace('/','-')
            #tags = article.select('a[href*="category"]')
            #tags = article.select('span.body')[2].select('a')
            #fb_iframe = article.select('iframe')[1]
            #response = urllib.request.urlopen(fb_iframe.attrs['src'])
            #iframe_soup = BeautifulSoup(response)
            share = article.select('span.shareCount')[0].text
            if share.find('K') != -1:
                share_num = float(share.split(' ')[0])*1000
                share_num = int(share_num)
            else:    
                share_num = share.split(' ')[0] #share = share.split(' '),share_num = share[0]

            #tag_string = ''
            #for tag in tags:
                #tag_string += tag.text + ','
            if date == today:
                #print(title)
                #print(tag_string)
                #print(type(share))
                row_cells = table.add_row().cells
                row_cells[0].text = title
                row_cells[1].text = str(share_num)
                row_cells[2].text = ''
                row_cells[3].text = 'TechOrange'
                
#                 sql = '''
#                     INSERT INTO `new_media`.`articles`(`brand`,`title`,`date`,`tags`,`share`) 
#                     VALUES ('{}','{}','{}','{}','{}')
#                     '''.format('TechOrange',title,date,'',share_num)
#                 #print(sql)
#                 cursor.execute(sql)
#                 connection.commit()
            
    document.save(today + '新媒體文章抓取.docx')    
    connection.close()
    driver.close()
except Exception as e:
    print(e)
    connection.close()
    driver.close()